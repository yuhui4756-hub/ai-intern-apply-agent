from __future__ import annotations

from pathlib import Path

from docx import Document


def read_resume_text(file_path: str) -> str:
    path = Path(file_path.strip('"')).expanduser()
    if not path.exists():
        return ""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return read_docx(path)
    if suffix == ".pdf":
        return read_pdf(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    return ""


def read_docx(path: Path) -> str:
    doc = Document(str(path))
    lines: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted and not reader.decrypt(""):
            return ""
        return "\n".join(text.strip() for page in reader.pages if (text := (page.extract_text() or "").strip()))
    except Exception:
        return ""
